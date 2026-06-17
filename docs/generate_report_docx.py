"""
Generates CloudDrive_Technical_Report.docx — the full project report as a Word
document: introduction, architecture, flow diagrams, data model, a technology
stack with WHY each tool was used and its BENEFIT, real screenshots,
deployment, testing, CI/CD, cost, and conclusion.

Diagrams are drawn with reportlab, rendered to PNG (renderPDF + pdftoppm), then
embedded. Run AFTER capture_screenshots.py + capture_shared.py.

Run:  python docs/generate_report_docx.py
"""
import os
import subprocess

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.graphics import renderPDF
# reuse the diagram builders + colors from the PDF generator
import generate_report as R

HERE = os.path.dirname(__file__)
IMG = os.path.join(HERE, "img")
OUT = os.path.join(HERE, "CloudDrive_Technical_Report.docx")
BLUE = RGBColor(0x25, 0x63, 0xEB)
INK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)


# ---------- diagram -> png ----------
def diagram_png(drawing, name, dpi=200):
    pdf = f"/tmp/{name}.pdf"
    renderPDF.drawToFile(drawing, pdf)
    subprocess.run(["pdftoppm", "-png", "-r", str(dpi), "-singlefile", pdf,
                    os.path.join(IMG, name)], check=True)
    return os.path.join(IMG, name + ".png")


# ---------- docx helpers ----------
def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE if level == 1 else INK
    return h


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


def caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)
    return p


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Medium Shading 1 Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        t.autofit = False
        t.allow_autofit = False
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = w
    return t


def image(doc, path, width_in, cap):
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def mono(doc, text):
    """A monospaced code/command block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    lines = text.strip("\n").split("\n")
    for i, ln in enumerate(lines):
        r = p.add_run(ln + ("\n" if i < len(lines) - 1 else ""))
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x0B, 0x30, 0x66)
    return p


# Section titles (H1) that appear in the Table of Contents
TOC_TITLES = [
    "1. Introduction", "2. Key Features", "3. System Architecture",
    "4. The Pre-signed URL Upload Flow", "5. Data Model",
    "6. Technology Stack — Why Each Tool & Its Benefit", "7. Security Model",
    "8. User Interface", "9. Containerization with Docker",
    "10. Cloud Storage on Amazon S3", "11. Cloud Deployment", "12. Testing",
    "13. CI/CD Pipeline", "14. Cost & Free Tier", "15. Conclusion & Future Work",
]


def toc_render(doc, entries):
    """entries: list of (title, page) or None (placeholder pass for pagination)."""
    p = doc.add_paragraph()
    r = p.add_run("Table of Contents")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(10)

    if entries is None:
        # pass 1 — reserve one page so pagination matches the final version
        doc.add_paragraph(" ")
        doc.add_page_break()
        return

    for title, page in entries:
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(4)
        # right-aligned tab with a dotted leader -> "Title ......... 3"
        line.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = line.add_run(title)
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK
        pg = line.add_run("\t" + str(page))
        pg.font.size = Pt(10.5)
        pg.font.color.rgb = INK
    doc.add_page_break()


def detect_toc_pages(pdf_path):
    """Find the page number each section heading lands on in the rendered PDF."""
    from pypdf import PdfReader
    reader = PdfReader(pdf_path)
    found = {}
    for i, page in enumerate(reader.pages, start=1):
        norm = " ".join((page.extract_text() or "").split())
        for t in TOC_TITLES:
            if t in found:
                continue
            key = t.split("—")[0].strip()  # tolerate em-dash extraction differences
            if key in norm:
                found[t] = i
    # fall back to sequential if any missed
    return [(t, found.get(t, "")) for t in TOC_TITLES]


def title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("☁ CloudDrive")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = BLUE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Report")
    r.font.size = Pt(20)
    r.font.color.rgb = INK
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Scalable File Storage on AWS — a Mini Google Drive\nFinal-Year Cloud Computing Project")
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Author: Shehriyar Ali Rustam\nRepository: github.com/Shehriyar-Ali-Rustam/clouddrive")
    r.font.size = Pt(11)
    doc.add_page_break()


# ---------- WHY/BENEFIT content for every tool ----------
STACK = [
    ("Backend & Language", [
        ("Python 3.12", "Readable, productive language with a huge ecosystem for web and AWS. Benefit: fast development and excellent library support (FastAPI, boto3)."),
        ("FastAPI", "Modern async web framework with built-in validation and auto-generated API docs. Benefit: high performance, less boilerplate, fewer bugs."),
        ("Uvicorn", "Lightning-fast ASGI server that runs the async FastAPI app. Benefit: handles many concurrent requests efficiently."),
        ("SQLAlchemy", "Database-agnostic ORM mapping Python objects to tables. Benefit: switch SQLite (local) to PostgreSQL (cloud) with no code change; avoids hand-written SQL."),
        ("Pydantic", "Validates request/response data and app settings using type hints. Benefit: bad input is rejected automatically; configuration is type-safe."),
    ]),
    ("Frontend", [
        ("HTML5 / CSS3", "Standard markup and styling for the dashboard UI. Benefit: lightweight, responsive, no build tooling required."),
        ("Vanilla JavaScript", "Implements the upload flow, drag-and-drop and sharing without a framework. Benefit: zero dependencies, fast to load, easy to explain in a viva."),
    ]),
    ("Security & Authentication", [
        ("JWT (python-jose)", "Stateless, signed bearer tokens for authentication. Benefit: no server-side sessions, so the API scales horizontally."),
        ("bcrypt", "Industry-standard one-way password hashing. Benefit: passwords stay safe even if the database is ever leaked."),
    ]),
    ("Storage, Database & SDK", [
        ("Amazon S3", "Cloud object storage for the actual file bytes. Benefit: virtually unlimited, 99.999999999% durable, cheap, and offloads heavy file traffic from the server."),
        ("PostgreSQL", "Robust relational database for metadata in production. Benefit: reliable, ACID-compliant, scales well."),
        ("SQLite", "Zero-configuration file database for local development. Benefit: run and test instantly with no database server."),
        ("psycopg2", "PostgreSQL driver for Python. Benefit: stable, well-supported connectivity to RDS."),
        ("boto3", "Official AWS SDK for Python. Benefit: generates S3 pre-signed URLs and manages buckets with a few lines of code."),
    ]),
    ("AWS Cloud Infrastructure", [
        ("AWS IAM", "Defines least-privilege roles and policies. Benefit: the app can touch only its own bucket — limits damage if compromised."),
        ("Amazon RDS", "Managed PostgreSQL service. Benefit: automated backups, patching and scaling — no manual database administration."),
        ("Amazon ECS Fargate", "Runs the API in serverless containers. Benefit: auto-scales 1→4 under load with no servers to manage."),
        ("Application Load Balancer", "Distributes traffic across API containers. Benefit: high availability and automatic health checks."),
        ("Amazon VPC", "Isolated virtual network with public/private subnets. Benefit: the database and app are not exposed to the public internet."),
        ("Amazon ECR", "Private registry for Docker images. Benefit: secure, integrates directly with ECS deployments."),
    ]),
    ("DevOps, IaC & Tooling", [
        ("Terraform", "Defines the whole AWS stack as code. Benefit: reproducible, version-controlled infrastructure built or destroyed with one command (controls cost)."),
        ("Docker", "Packages the app and its dependencies into a container. Benefit: runs identically everywhere — no 'works on my machine' issues."),
        ("GitHub Actions", "Automated CI/CD pipeline. Benefit: every push runs the tests and can auto-deploy — bugs are caught early."),
        ("Git & GitHub", "Version control and code hosting. Benefit: full history, backup, and collaboration."),
        ("pytest", "Automated testing framework (19 tests). Benefit: confidence that new changes don't break existing features."),
    ]),
    ("Hosting & Networking", [
        ("Render", "Platform-as-a-Service for permanent hosting. Benefit: a permanent HTTPS URL plus a free managed PostgreSQL, deployed straight from GitHub."),
        ("Cloudflare Tunnel", "Exposes the local app on a public HTTPS URL. Benefit: a free, shareable link for live demos with no deployment."),
    ]),
]


def build(out=OUT, toc_entries=None):
    doc = Document()
    doc.core_properties.title = "CloudDrive — Technical Report"
    doc.core_properties.author = "Shehriyar Ali Rustam"

    # default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    title_page(doc)
    toc_render(doc, toc_entries)

    # 1. Introduction
    heading(doc, "1. Introduction")
    heading(doc, "1.1 Problem Statement", 2)
    body(doc, "Storing and sharing files across devices and people is a universal need. Traditional "
              "approaches that stream every file through an application server do not scale: the "
              "server becomes a bottleneck and a single point of failure. CloudDrive solves this with "
              "cloud object storage and pre-signed URLs, keeping the application server stateless and "
              "horizontally scalable.")
    heading(doc, "1.2 Objectives", 2)
    for t in ["Provide secure user accounts with authentication and per-user isolation.",
              "Allow uploading, downloading, organising and deleting files in the cloud.",
              "Enable sharing — with registered users and via public expiring links.",
              "Demonstrate a scalable, cloud-native architecture (S3, RDS, ECS, IAM).",
              "Automate infrastructure (Terraform) and delivery (CI/CD)."]:
        bullet(doc, t)
    heading(doc, "1.3 Scope", 2)
    body(doc, "The system covers authentication, file/folder management, quotas, sharing, and a "
              "polished web UI. It runs in two interchangeable modes — local (SQLite + disk) for "
              "offline development and AWS (S3 + PostgreSQL) for the cloud — selected purely by "
              "configuration, with no code changes.")

    # 2. Features
    heading(doc, "2. Key Features")
    add_table(doc, ["Feature", "Description"], [
        ["Authentication", "Email + password sign-up/login, JWT tokens, bcrypt-hashed passwords"],
        ["File upload/download", "Browser transfers bytes directly to storage via pre-signed URLs"],
        ["Organisation", "Files and folders, isolated per user"],
        ["Sharing", "Share with a registered user, or create a public expiring link"],
        ["Storage quota", "Per-user limit with a live usage bar"],
        ["Versioning", "S3 bucket versioning keeps previous file versions"],
        ["Two run modes", "Local (SQLite+disk) and Cloud (S3+PostgreSQL) via one setting"],
    ], widths=[Inches(1.8), Inches(4.5)])

    # 3. Architecture
    heading(doc, "3. System Architecture")
    body(doc, "The design separates concerns: the database stores only metadata (who owns what, file "
              "names, sizes, storage keys, share rules) while Amazon S3 stores the actual file bytes. "
              "The browser talks directly to S3 for the heavy data transfer.")
    image(doc, diagram_png(R.architecture_diagram(), "diag_arch"), 5.7,
          "Figure 1: High-level cloud architecture.")
    heading(doc, "3.1 Components", 2)
    add_table(doc, ["Component", "Technology", "Responsibility"], [
        ["Frontend", "HTML/CSS/JS SPA", "Dashboard UI, upload flow, sharing"],
        ["API", "FastAPI on ECS Fargate", "Auth, metadata, pre-signed URL issuing"],
        ["File storage", "Amazon S3", "Stores file bytes; versioned & encrypted"],
        ["Database", "PostgreSQL / SQLite", "Users, files, folders, shares"],
        ["Load balancer", "Application Load Balancer", "Distributes traffic, health checks"],
        ["Auth", "JWT + bcrypt", "Stateless authentication"],
        ["IaC", "Terraform", "Provisions the entire AWS stack"],
    ], widths=[Inches(1.3), Inches(2.0), Inches(3.0)])

    # 4. Upload flow
    heading(doc, "4. The Pre-signed URL Upload Flow")
    body(doc, "This is the central cloud concept. The API never streams file bytes; it issues a "
              "short-lived signed permission slip (a pre-signed URL) and the browser uploads directly "
              "to S3. This keeps the API lightweight and horizontally scalable.")
    flow = R.pipeline([
        ("1. Init", ["Browser tells API", "name + size"]),
        ("2. Ticket", ["API checks quota,", "returns pre-signed URL"]),
        ("3. Upload", ["Browser PUTs bytes", "directly to S3"]),
        ("4. Complete", ["API marks done,", "updates quota"]),
    ])
    image(doc, diagram_png(flow, "diag_upload"), 6.2, "Figure 2: Three-step pre-signed upload sequence.")

    # 5. Data model
    heading(doc, "5. Data Model")
    body(doc, "Four tables capture all metadata. File bytes live in S3, referenced by a storage key.")
    add_table(doc, ["Table", "Key Fields", "Purpose"], [
        ["users", "id, email, password_hash, storage_used, storage_quota", "Accounts & quota"],
        ["folders", "id, name, owner_id, parent_id", "Folder hierarchy"],
        ["files", "id, name, s3_key, size, owner_id, folder_id, uploaded", "File metadata"],
        ["shares", "id, file_id, shared_with_user_id, public_token, expires_at", "Sharing rules"],
    ], widths=[Inches(0.9), Inches(3.4), Inches(2.0)])

    # 6. Tech stack with WHY / BENEFIT
    heading(doc, "6. Technology Stack — Why Each Tool & Its Benefit")
    body(doc, "Every technology below was chosen for a specific reason. The table for each layer "
              "explains why it was used and the benefit it brings to CloudDrive.")
    for cat, items in STACK:
        heading(doc, cat, 2)
        add_table(doc, ["Tool / Service", "Why it was used & its benefit"],
                  [[name, why] for name, why in items],
                  widths=[Inches(1.7), Inches(4.6)])

    # 7. Security
    heading(doc, "7. Security Model")
    add_table(doc, ["Control", "How it protects the system"], [
        ["Password hashing", "bcrypt one-way hashes; plaintext never stored"],
        ["JWT auth", "Signed, expiring bearer tokens on every API call"],
        ["Pre-signed URLs", "Time-limited (15 min), scoped to a single object"],
        ["Private bucket", "S3 blocks all public access; encrypted at rest (AES-256)"],
        ["IAM least-privilege", "API role may touch only this one bucket"],
        ["Per-user isolation", "Every query filtered by owner_id"],
        ["Quota enforcement", "Prevents a single user exhausting storage"],
    ], widths=[Inches(1.8), Inches(4.5)])

    # 8. UI screenshots
    heading(doc, "8. User Interface")
    body(doc, "Real screenshots of the running application (the badge shows storage: S3).")
    image(doc, os.path.join(IMG, "01_login.png"), 4.6, "Figure 3: Sign-up / login screen.")
    image(doc, os.path.join(IMG, "02_dashboard.png"), 6.2, "Figure 4: Dashboard — files on AWS S3, quota bar, type icons.")
    image(doc, os.path.join(IMG, "03_share_modal.png"), 6.2, "Figure 5: Share dialog — share with a user or create a public link.")
    image(doc, os.path.join(IMG, "04_public_link.png"), 6.2, "Figure 6: A generated public expiring link.")
    image(doc, os.path.join(IMG, "05_shared_with_me.png"), 6.2, "Figure 7: 'Shared with me' — files another user shared.")

    # 9. Containerization with Docker
    heading(doc, "9. Containerization with Docker")
    body(doc, "The application is packaged as a Docker image so it runs identically on any machine "
              "and in the cloud. The Dockerfile installs the dependencies, copies the code, and starts "
              "the API server:")
    mono(doc,
         "FROM python:3.12-slim\n"
         "WORKDIR /app\n"
         "COPY requirements.txt .\n"
         "RUN pip install --no-cache-dir -r requirements.txt\n"
         "COPY backend ./backend\n"
         "COPY frontend ./frontend\n"
         "EXPOSE 8000\n"
         'CMD ["uvicorn", "backend.main:app", "--host", "0.0.0.0", "--port", "8000"]')
    body(doc, "The image was built and run as a live container serving the app on real AWS S3:")
    mono(doc,
         "$ docker build -t clouddrive-api:latest .\n"
         "$ docker run -d --name clouddrive -p 8001:8000 \\\n"
         "      -e STORAGE_BACKEND=s3 -e S3_BUCKET=clouddrive-066496974419 \\\n"
         "      -v ~/.aws:/root/.aws:ro clouddrive-api:latest")
    image(doc, os.path.join(IMG, "docker.png"), 6.2,
          "Figure 8: Docker image built (395 MB) and the container running, serving the app on S3.")
    body(doc, "This same image is pushed to Amazon ECR and run on ECS Fargate in production — locally "
              "one container runs the app; in the cloud ECS runs several copies behind the load "
              "balancer and autoscales them.")

    # 10. Cloud Storage on S3
    heading(doc, "10. Cloud Storage on Amazon S3")
    body(doc, "File bytes are stored in a real Amazon S3 bucket, uploaded directly by the browser via "
              "pre-signed URLs. The bucket has versioning enabled and AES-256 server-side encryption, "
              "and blocks all public access. The screenshot below verifies the live bucket contents "
              "and settings via the AWS CLI:")
    image(doc, os.path.join(IMG, "s3_bucket.png"), 6.2,
          "Figure 9: Real objects stored in the S3 bucket, with versioning enabled and AES-256 encryption.")

    # 11. Deployment
    heading(doc, "11. Cloud Deployment")
    body(doc, "The application can be run or deployed three ways, all sharing the same S3 storage:")
    dep = R.pipeline([
        ("Local", ["SQLite + disk", "or S3 mode"]),
        ("Render", ["Permanent URL", "+ managed Postgres"]),
        ("AWS ECS", ["Docker image,", "autoscaling (Terraform)"]),
    ], box_h=46)
    image(doc, diagram_png(dep, "diag_deploy"), 6.0, "Figure 10: Deployment options (all use Amazon S3 for files).")
    body(doc, "For demonstrations, a Cloudflare tunnel exposes the local app on a public HTTPS URL so "
              "it can be opened from any device. The full AWS stack is defined in Terraform (validated) "
              "and deploys via the CI/CD pipeline below.")

    # 12. Testing
    heading(doc, "12. Testing")
    body(doc, "Nineteen automated tests (pytest) cover the critical paths and security:")
    add_table(doc, ["Area", "Examples", "Tests"], [
        ["Authentication", "signup, login, wrong password, duplicate email", "6"],
        ["Files", "upload lifecycle, listing, quota, delete, isolation", "7"],
        ["Sharing & security", "user share, public link, tampered/expired URL", "6"],
    ], widths=[Inches(1.6), Inches(3.6), Inches(1.1)])
    body(doc, "All 19 tests pass and run automatically on every push via GitHub Actions.")

    # 13. CI/CD
    heading(doc, "13. CI/CD Pipeline")
    body(doc, "Two GitHub Actions workflows automate delivery: CI runs the tests on every push, and "
              "Deploy builds the Docker image, pushes it to Amazon ECR, and rolls it out on ECS.")
    cicd = R.pipeline([
        ("git push", ["Code to GitHub"]),
        ("CI: Tests", ["Run all 19 tests"]),
        ("Build", ["Docker image", "-> Amazon ECR"]),
        ("Deploy", ["Rolling update", "on ECS"]),
    ], box_h=46)
    image(doc, diagram_png(cicd, "diag_cicd"), 6.2, "Figure 11: Continuous integration & deployment flow.")
    image(doc, os.path.join(IMG, "gh_actions.png"), 6.2,
          "Figure 12: GitHub Actions — CI and Deploy workflows running green on every push.")
    image(doc, os.path.join(IMG, "gh_actions_run.png"), 6.2,
          "Figure 13: A CI run executing the automated test suite.")

    # 14. Cost
    heading(doc, "14. Cost & Free Tier")
    add_table(doc, ["Service", "Cost note"], [
        ["Amazon S3", "Free tier: 5 GB + 20k GETs + 2k PUTs — demo cost ~ $0"],
        ["IAM / ECR", "Free"],
        ["RDS (db.t3.micro)", "Free-tier eligible for 12 months"],
        ["ALB / NAT / ECS", "Billed hourly — destroyed after demos via 'terraform destroy'"],
        ["Render / Tunnel", "Free tiers"],
    ], widths=[Inches(1.9), Inches(4.4)])

    # 15. Conclusion
    heading(doc, "15. Conclusion & Future Work")
    body(doc, "CloudDrive delivers a complete, working and tested cloud application that demonstrates "
              "the key principles of cloud computing — object storage, pre-signed URLs, managed "
              "databases, stateless scalable services, least-privilege security, infrastructure as "
              "code, and automated delivery. It runs on real AWS S3 and is fully reproducible from "
              "version control.")
    for t in ["Future: multipart upload for very large files (S3 multipart + Lambda confirm).",
              "Future: file previews and full-text search over metadata.",
              "Future: organisation/team workspaces and granular permissions."]:
        bullet(doc, t)
    body(doc, "Appendix — Run locally: bash run.sh  ·  Run tests: bash run_tests.sh  ·  "
              "Public demo: bash demo.sh")

    doc.save(out)
    return out


def main():
    # Pass 1: build with a placeholder TOC, render to PDF, find heading pages.
    tmp = "/tmp/_cd_report_pass1.docx"
    build(out=tmp, toc_entries=None)
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", "/tmp", tmp], check=True,
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    entries = detect_toc_pages("/tmp/_cd_report_pass1.pdf")
    print("TOC pages:", entries)
    # Pass 2: build the final doc with the real page numbers.
    build(out=OUT, toc_entries=entries)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
