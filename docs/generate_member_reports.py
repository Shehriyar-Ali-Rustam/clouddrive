"""
Generates 3 per-member reports from the same CloudDrive project, one per team
member / feature module. The collective report (generate_report_docx.py) is
left untouched; this reuses its helpers, styling, diagrams and screenshots.

  Member 1  -> Authentication & User Accounts   (+ Security/IAM/CI)
  Member 2  -> File Storage & Upload/Download    (+ S3/Terraform/Docker)  [Shehriyar]
  Member 3  -> Sharing & Collaboration           (+ Deployment/Networking)

Run:  python docs/generate_member_reports.py
"""
import os
import subprocess

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Reuse everything from the collective generator (helpers, STACK, colours, diagrams)
import generate_report_docx as G

HERE = os.path.dirname(__file__)
IMG = G.IMG

# tool -> "why & benefit" rationale, reused from the collective report's STACK
TOOL_WHY = {name: why for _cat, items in G.STACK for name, why in items}


def member_title_page(doc, module_title, member_name):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("☁ CloudDrive"); r.bold = True; r.font.size = Pt(38); r.font.color.rgb = G.BLUE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(module_title); r.font.size = Pt(19); r.font.color.rgb = G.INK
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Team Module Report"); r.font.size = Pt(12); r.font.color.rgb = G.GREY
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Prepared by: {member_name}\n"
                  "Project: CloudDrive — Scalable File Storage on AWS\n"
                  "Repository: github.com/Shehriyar-Ali-Rustam/clouddrive")
    r.font.size = Pt(11)
    doc.add_page_break()


def section_titles(cfg):
    """The H1 headings (in order) that go in this report's Table of Contents."""
    return [
        "1. Project Overview",
        f"2. My Module: {cfg['module_title']}",
        cfg["flow_title"],
        "4. Data & Code",
        "5. Technologies I Used — Why & Benefit",
        cfg["cloud_heading"],
        "7. Testing",
        "8. Conclusion & My Contribution",
    ]


def detect_pages(pdf_path, titles):
    from pypdf import PdfReader
    reader = PdfReader(pdf_path)
    found = {}
    for i, page in enumerate(reader.pages, start=1):
        norm = " ".join((page.extract_text() or "").split())
        for t in titles:
            if t in found:
                continue
            key = t.split("—")[0].strip()  # tolerate em-dash extraction
            if key in norm:
                found[t] = i
    return [(t, found.get(t, "")) for t in titles]


TEAM_ROWS = [
    ["Authentication & User Accounts", "Member 1"],
    ["File Storage & Upload/Download", "Shehriyar Ali Rustam"],
    ["Sharing & Collaboration", "Member 3"],
]


def build_member(cfg, out, toc_entries):
    doc = Document()
    doc.core_properties.title = f"CloudDrive — {cfg['module_title']}"
    doc.core_properties.author = cfg["member_name"]
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    member_title_page(doc, cfg["module_title"], cfg["member_name"])
    G.toc_render(doc, toc_entries)

    # 1. Project Overview (shared context + team split)
    G.heading(doc, "1. Project Overview")
    G.body(doc, "CloudDrive is a cloud-based file-storage web application — a simplified Google Drive — "
                "that lets users securely upload, organise, download and share files on AWS. It was built "
                "by a team of three, each owning one feature module (and a slice of the cloud/DevOps work). "
                "This report covers my module.")
    rows = [[m, n + (" (this report)" if n == cfg["member_name"] else "")] for m, n in TEAM_ROWS]
    G.add_table(doc, ["Project module", "Team member"], rows, widths=[Inches(3.5), Inches(2.8)])

    # 2. My Module
    G.heading(doc, f"2. My Module: {cfg['module_title']}")
    G.body(doc, cfg["feature_intro"])
    for b in cfg["feature_bullets"]:
        G.bullet(doc, b)
    for img, w, cap in cfg.get("feature_images", []):
        G.image(doc, os.path.join(IMG, img), w, cap)

    # 3. How It Works (flow diagram)
    G.heading(doc, cfg["flow_title"])
    G.body(doc, cfg["flow_intro"])
    flow = G.R.pipeline(cfg["flow_steps"], box_h=cfg.get("flow_box_h", 50))
    G.image(doc, G.diagram_png(flow, f"flow_{cfg['key']}"), 6.3, cfg["flow_caption"])

    # 4. Data & Code
    G.heading(doc, "4. Data & Code")
    G.add_table(doc, cfg["model_headers"], cfg["model_rows"], widths=[Inches(2.3), Inches(4.0)])
    G.body(doc, "Key endpoints / entry points:")
    G.mono(doc, cfg["endpoints"])
    G.body(doc, cfg["code_note"])

    # 5. Technologies I used — why & benefit
    G.heading(doc, "5. Technologies I Used — Why & Benefit")
    G.add_table(doc, ["Tool / Service", "Why it was used & its benefit"],
                [[t, TOOL_WHY[t]] for t in cfg["tools"]],
                widths=[Inches(1.7), Inches(4.6)])

    # 6. Cloud / DevOps contribution
    G.heading(doc, cfg["cloud_heading"])
    G.body(doc, cfg["cloud_body"])
    for b in cfg["cloud_bullets"]:
        G.bullet(doc, b)
    if cfg.get("cloud_table_rows"):
        G.add_table(doc, cfg["cloud_table_headers"], cfg["cloud_table_rows"],
                    widths=[Inches(2.0), Inches(4.3)])
    if cfg.get("dockerfile"):
        G.mono(doc, cfg["dockerfile"])
    if cfg.get("cloud_diagram_steps"):
        cd = G.R.pipeline(cfg["cloud_diagram_steps"], box_h=46)
        G.image(doc, G.diagram_png(cd, f"clouddiag_{cfg['key']}"), 6.0, cfg["cloud_diagram_caption"])
    for img, w, cap in cfg.get("cloud_images", []):
        G.image(doc, os.path.join(IMG, img), w, cap)

    # 7. Testing
    G.heading(doc, "7. Testing")
    G.add_table(doc, cfg["test_headers"], cfg["test_rows"], widths=[Inches(2.6), Inches(3.7)])
    G.body(doc, cfg["test_note"])

    # 8. Conclusion
    G.heading(doc, "8. Conclusion & My Contribution")
    G.body(doc, cfg["conclusion"])

    doc.save(out)
    return out


# ====================================================================== #
#  THE THREE MEMBER CONFIGS
# ====================================================================== #
MEMBERS = [
    {
        "key": "auth",
        "member_name": "Member 1",
        "module_title": "Authentication & User Accounts",
        "out": os.path.join(HERE, "CloudDrive_Report_1_Authentication.docx"),
        "feature_intro":
            "This module secures the application: it lets users create accounts, log in, and proves "
            "their identity on every request. It also enforces per-user storage quotas and guarantees "
            "that each user can only ever see their own data.",
        "feature_bullets": [
            "Sign-up and login with email and password.",
            "Passwords stored only as bcrypt one-way hashes — never in plain text.",
            "Stateless authentication using signed JSON Web Tokens (JWT).",
            "A reusable 'current user' guard that protects every private endpoint.",
            "Per-user storage quota shown as a live usage bar.",
            "Strict per-user isolation — every query is filtered by the owner.",
        ],
        "feature_images": [("01_login.png", 4.6, "Figure 1: The sign-up / login screen.")],
        "flow_title": "3. How It Works — Authentication Flow",
        "flow_intro":
            "Authentication is stateless: the server hands the browser a signed token at login, and the "
            "browser presents it on every later request. No server-side sessions are kept.",
        "flow_steps": [
            ("1. Login", ["email + password", "sent to API"]),
            ("2. Verify", ["bcrypt checks", "the password"]),
            ("3. Issue JWT", ["signed token", "returned"]),
            ("4. Authorize", ["Bearer token on", "every request"]),
        ],
        "flow_caption": "Figure 2: Stateless JWT authentication flow.",
        "model_headers": ["users table — field", "Purpose"],
        "model_rows": [
            ["email", "unique login identifier"],
            ["password_hash", "bcrypt hash of the password (never plaintext)"],
            ["storage_used / storage_quota", "bytes used vs allowed (the quota)"],
            ["created_at", "account creation time"],
        ],
        "endpoints":
            "POST /api/auth/signup   -> create account, return JWT\n"
            "POST /api/auth/login    -> verify password, return JWT\n"
            "GET  /api/auth/me       -> current user + quota (guarded)",
        "code_note":
            "Implemented in backend/auth.py (hashing, JWT, the get_current_user guard) and the auth "
            "endpoints in backend/main.py; the User model lives in backend/models.py. Frontend logic: "
            "submitAuth / enterApp / updateQuota / logout in frontend/app.js.",
        "tools": ["Python 3.12", "FastAPI", "JWT (python-jose)", "bcrypt", "Pydantic",
                  "AWS IAM", "GitHub Actions", "pytest"],
        "cloud_heading": "6. My Cloud & DevOps Contribution — Security, IAM & CI",
        "cloud_body":
            "Beyond the feature, I own the project's security posture and its continuous-integration "
            "testing.",
        "cloud_bullets": [
            "Security model: bcrypt password hashing, expiring JWTs, and per-user isolation.",
            "AWS IAM least-privilege roles so the application can access only its own S3 bucket.",
            "GitHub Actions CI that runs the entire test suite automatically on every push.",
        ],
        "cloud_table_headers": ["Security control", "How it protects the system"],
        "cloud_table_rows": [
            ["bcrypt hashing", "passwords stay safe even if the database leaks"],
            ["JWT (expiring)", "stateless auth, so the API scales horizontally"],
            ["IAM least-privilege", "the app is limited to a single S3 bucket"],
            ["Per-user isolation", "users can never see each other's data"],
        ],
        "cloud_images": [("gh_actions_run.png", 6.2,
                          "Figure 3: GitHub Actions CI running the automated test suite.")],
        "test_headers": ["Test (tests/test_auth.py)", "What it checks"],
        "test_rows": [
            ["signup returns token", "account creation issues a JWT"],
            ["duplicate email rejected", "no two accounts share an email (400)"],
            ["login wrong password", "rejected with 401"],
            ["me requires auth", "protected endpoint blocks anonymous access"],
            ["me returns user + quota", "identity and quota returned correctly"],
        ],
        "test_note": "6 automated tests cover this module (tests/test_auth.py), all passing in CI.",
        "conclusion":
            "I delivered the authentication and account system — secure signup/login with JWT and bcrypt, "
            "the guard that protects every endpoint, and per-user quotas and isolation — plus the "
            "project's IAM security posture and the GitHub Actions CI that keeps the whole codebase tested.",
    },
    {
        "key": "storage",
        "member_name": "Shehriyar Ali Rustam",
        "module_title": "File Storage & Upload/Download",
        "out": os.path.join(HERE, "CloudDrive_Report_2_FileStorage.docx"),
        "feature_intro":
            "This module is the heart of CloudDrive: it stores users' files in the cloud and moves bytes "
            "efficiently using pre-signed URLs, so the API server never becomes a bottleneck.",
        "feature_bullets": [
            "Upload, download, list and delete files; organise them in folders.",
            "Files are stored as objects in Amazon S3 — not on the application server.",
            "Pre-signed URLs let the browser transfer bytes directly to and from S3.",
            "A pluggable storage layer: real S3 in the cloud, local disk for offline development — same code.",
            "Quota is checked before an upload, billed after it, and refunded on delete.",
        ],
        "feature_images": [("02_dashboard.png", 6.2,
                            "Figure 1: The file dashboard (the badge shows storage: S3).")],
        "flow_title": "3. How It Works — The Pre-signed Upload Flow",
        "flow_intro":
            "The API never streams file bytes. It issues a short-lived signed permission slip (a "
            "pre-signed URL) and the browser uploads straight to S3 — the single most important "
            "decision for scalability.",
        "flow_steps": [
            ("1. Init", ["Browser tells API", "name + size"]),
            ("2. Ticket", ["API checks quota,", "returns pre-signed URL"]),
            ("3. Upload", ["Browser PUTs bytes", "directly to S3"]),
            ("4. Complete", ["API marks done,", "updates quota"]),
        ],
        "flow_caption": "Figure 2: Three-step pre-signed upload — bytes never pass through the API.",
        "model_headers": ["files table — field", "Purpose"],
        "model_rows": [
            ["s3_key", "the object's location in S3"],
            ["size / content_type", "byte size and MIME type"],
            ["owner_id / folder_id", "the owner and the containing folder"],
            ["uploaded", "flipped true once the bytes have landed"],
        ],
        "endpoints":
            "POST   /api/files/init           -> pre-signed upload URL\n"
            "POST   /api/files/{id}/complete  -> confirm upload, bill quota\n"
            "GET    /api/files                -> list my files\n"
            "GET    /api/files/{id}/download  -> pre-signed download URL\n"
            "DELETE /api/files/{id}           -> delete + refund quota\n"
            "POST/GET /api/folders            -> folder management",
        "code_note":
            "Implemented in backend/storage.py (the S3/local storage abstraction) and the file/folder "
            "endpoints in backend/main.py; the File and Folder models in backend/models.py. Frontend: "
            "uploadOne / putWithProgress / loadFiles / renderGrid in frontend/app.js.",
        "tools": ["Python 3.12", "FastAPI", "SQLAlchemy", "Amazon S3", "boto3", "SQLite",
                  "Terraform", "Docker"],
        "cloud_heading": "6. My Cloud & DevOps Contribution — S3, Terraform & Docker",
        "cloud_body": "I own the project's cloud storage and its packaging and infrastructure-as-code.",
        "cloud_bullets": [
            "An Amazon S3 bucket with versioning and AES-256 encryption, with all public access blocked.",
            "Terraform definitions that provision the S3 bucket and supporting infrastructure as code.",
            "A Dockerfile that containerizes the whole app; the same image runs locally and on ECS.",
        ],
        "dockerfile":
            "FROM python:3.12-slim\n"
            "WORKDIR /app\n"
            "RUN pip install --no-cache-dir -r requirements.txt\n"
            "COPY backend ./backend ; COPY frontend ./frontend\n"
            'CMD ["uvicorn", "backend.main:app", "--host", "0.0.0.0", "--port", "8000"]',
        "cloud_images": [
            ("s3_bucket.png", 6.2, "Figure 3: Real objects in the S3 bucket (versioning + AES-256)."),
            ("docker.png", 6.2, "Figure 4: The app running inside a Docker container, serving S3."),
        ],
        "test_headers": ["Test (tests/test_files.py)", "What it checks"],
        "test_rows": [
            ["upload lifecycle", "init -> PUT -> complete succeeds"],
            ["file appears in list", "an uploaded file is listed"],
            ["download returns bytes", "downloaded content matches the original"],
            ["quota billed / freed", "quota updates on upload and on delete"],
            ["quota exceeded -> 413", "an over-quota upload is rejected"],
            ["cross-user blocked", "a user cannot access another's file"],
        ],
        "test_note": "7 automated tests cover this module (tests/test_files.py), all passing in CI.",
        "conclusion":
            "I delivered the cloud storage core — the pre-signed upload/download flow on Amazon S3, the "
            "pluggable storage layer, and folder and quota management — plus the Terraform and Docker "
            "packaging that make the application deployable to the cloud.",
    },
    {
        "key": "sharing",
        "member_name": "Member 3",
        "module_title": "Sharing & Collaboration",
        "out": os.path.join(HERE, "CloudDrive_Report_3_Sharing.docx"),
        "feature_intro":
            "This module lets users share their files — either with another registered user, or with "
            "anyone via a public link that can expire — and shows each user the files that others have "
            "shared with them.",
        "feature_bullets": [
            "Share a file with a registered user by email.",
            "Generate a public link that anyone can open, with an optional expiry time.",
            "A 'Shared with me' view listing files that others have shared.",
            "Secure downloads: links are short-lived pre-signed URLs scoped to a single file.",
            "Public links redirect straight to storage, so the server never streams the bytes.",
        ],
        "feature_images": [
            ("03_share_modal.png", 6.2, "Figure 1: The share dialog — share with a user or create a link."),
            ("04_public_link.png", 6.2, "Figure 2: A generated public expiring link."),
            ("05_shared_with_me.png", 6.0, "Figure 3: 'Shared with me' — files another user shared."),
        ],
        "flow_title": "3. How It Works — Sharing Flow",
        "flow_intro":
            "A share is recorded as a small database row — either pointing at a recipient user or holding "
            "an unguessable public token. Downloads always go through a short-lived pre-signed redirect.",
        "flow_steps": [
            ("1. Share", ["owner shares file", "with user or public"]),
            ("2. Record", ["Share row: user id", "or public token"]),
            ("3. Open", ["recipient opens link", "or 'Shared with me'"]),
            ("4. Download", ["pre-signed redirect", "to S3 (time-limited)"]),
        ],
        "flow_caption": "Figure 4: Sharing and secure download flow.",
        "model_headers": ["shares table — field", "Purpose"],
        "model_rows": [
            ["file_id", "the file being shared"],
            ["shared_with_user_id", "the recipient (for a user-to-user share)"],
            ["public_token", "unguessable token (for a public link)"],
            ["expires_at", "optional expiry time after which the link fails"],
        ],
        "endpoints":
            "POST /api/shares                     -> share with user OR create public link\n"
            "GET  /api/shared-with-me             -> files shared with me\n"
            "GET  /api/files/{id}/shared-download -> download a file shared with me\n"
            "GET  /api/public/{token}             -> public download (redirects to S3)",
        "code_note":
            "Implemented in the sharing endpoints of backend/main.py (create_share, shared_with_me, "
            "shared_download, public_download); the Share model in backend/models.py. Frontend: "
            "openShare / shareWithUser / makePublicLink / loadSharedWithMe in frontend/app.js.",
        "tools": ["Python 3.12", "FastAPI", "SQLAlchemy", "Render", "Cloudflare Tunnel",
                  "Amazon ECS Fargate", "Application Load Balancer", "Amazon VPC"],
        "cloud_heading": "6. My Cloud & DevOps Contribution — Deployment & Networking",
        "cloud_body": "I own getting the application onto the internet and the network that supports it.",
        "cloud_bullets": [
            "Deployment to Render (a permanent URL plus a managed PostgreSQL) and a Cloudflare tunnel for demos.",
            "The AWS network: a VPC with public/private subnets and an Application Load Balancer in front of ECS.",
            "The continuous-deployment workflow that builds the image and ships it.",
        ],
        "cloud_diagram_steps": [
            ("Local", ["dev / demo"]),
            ("Render", ["permanent URL", "+ Postgres"]),
            ("AWS ECS", ["containers", "behind ALB"]),
        ],
        "cloud_diagram_caption": "Figure 5: Deployment options (all use Amazon S3 for files).",
        "cloud_images": [("gh_actions.png", 6.2,
                          "Figure 6: GitHub Actions — CI and Deploy workflows running green.")],
        "test_headers": ["Test (tests/test_sharing_security.py)", "What it checks"],
        "test_rows": [
            ["share with user", "creates a share for that user"],
            ["unknown user -> 404", "cannot share to a non-existent account"],
            ["public link download", "the token resolves and downloads"],
            ["unknown token -> 404", "bad links are rejected"],
            ["tampered signature -> 403", "forged pre-signed URLs are blocked"],
        ],
        "test_note": "6 automated tests cover this module (tests/test_sharing_security.py).",
        "conclusion":
            "I delivered the sharing and collaboration features — user-to-user sharing, public expiring "
            "links, and the 'Shared with me' view with secure pre-signed downloads — plus the project's "
            "deployment (Render / ECS / tunnel) and the network (VPC and load balancer) that put it online.",
    },
]


def main():
    for cfg in MEMBERS:
        titles = section_titles(cfg)
        tmp = f"/tmp/_member_{cfg['key']}.docx"
        build_member(cfg, tmp, None)               # pass 1: placeholder TOC
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", "/tmp", tmp],
                       check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        entries = detect_pages(f"/tmp/_member_{cfg['key']}.pdf", titles)
        build_member(cfg, cfg["out"], entries)     # pass 2: real page numbers
        print("wrote", cfg["out"], "| TOC:", [(t.split('.')[0], p) for t, p in entries])


if __name__ == "__main__":
    main()
